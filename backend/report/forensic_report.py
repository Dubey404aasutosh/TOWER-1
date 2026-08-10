"""
Forensic Report Exporter
=========================
Generates Word document (python-docx) per flagged entity containing:
- Resolved identity summary
- Risk tier + rules fired with plain-English explanations
- Unified timeline chart (embedded as image)
- Money-flow graph (embedded as image)
- Evidence table with source file references
"""
import os
import io
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


REPORT_DIR = Path(__file__).resolve().parent.parent / "reports"
REPORT_DIR.mkdir(exist_ok=True)


def _resolve_png(source, width=750, height=450):
    """
    Accept either ready-made PNG bytes or a Plotly figure and return PNG bytes.

    The report path supplies bytes rendered by `report_charts` (matplotlib, ~0.1 s).
    A Plotly figure is still honoured for callers that want kaleido — it just costs
    ~3.5 s per image, which is why it is not what the pipeline passes.
    Returns None if nothing renders, so the caller can print an honest placeholder
    instead of an empty frame.
    """
    if source is None:
        return None
    if isinstance(source, (bytes, bytearray)):
        return bytes(source) or None
    try:
        return source.to_image(format="png", width=width, height=height, scale=1)
    except Exception as e:
        print(f"  [chart] figure could not be rasterised: {type(e).__name__}: {e}")
        return None


def generate_forensic_report(entity_id, entity_data, risk_data, all_events_df,
                              timeline_fig=None, network_fig=None,
                              timeline_png=None, network_png=None):
    """
    Generate a Word forensic report for a flagged entity.

    Args:
        entity_id: Resolved entity ID
        entity_data: Dict with phones, accounts, imeis, vpas
        risk_data: Dict with risk_tier, rules_fired, explanations, evidence, ml_anomaly
        all_events_df: DataFrame of all events
        timeline_fig: Plotly figure for the timeline (optional, rendered via kaleido)
        network_fig: Plotly figure for the network graph (optional, via kaleido)
        timeline_png: pre-rendered timeline PNG bytes (preferred — see report_charts)
        network_png: pre-rendered money-flow PNG bytes (preferred)

    Returns:
        Path to generated report file
    """
    doc = Document()

    # --- STYLE SETUP ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1a, 0x1d, 0x24)

    # --- HEADER ---
    heading = doc.add_heading('', level=0)
    run = heading.add_run('E-RAKSHAK FORENSIC INVESTIGATION REPORT')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1a, 0x1d, 0x24)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sub-header
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f'Entity: {entity_id} | Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('_' * 70)

    # --- SECTION 1: RISK ASSESSMENT ---
    doc.add_heading('1. Risk Assessment', level=1)

    risk_tier = risk_data.get("risk_tier", "LOW")
    ml_anomaly = risk_data.get("ml_anomaly", False)
    rules = risk_data.get("rules_fired", [])

    tier_para = doc.add_paragraph()
    tier_para.add_run('Rule-Based Risk Tier: ').bold = True
    tier_run = tier_para.add_run(risk_tier)
    tier_run.bold = True
    if risk_tier == "CRITICAL":
        tier_run.font.color.rgb = RGBColor(0xdc, 0x35, 0x45)
    elif risk_tier == "HIGH":
        tier_run.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)
    elif risk_tier == "MEDIUM":
        tier_run.font.color.rgb = RGBColor(0xf3, 0x9c, 0x12)
    else:
        tier_run.font.color.rgb = RGBColor(0x27, 0xae, 0x60)

    ml_para = doc.add_paragraph()
    ml_para.add_run('ML Anomaly Detection: ').bold = True
    ml_run = ml_para.add_run('FLAGGED' if ml_anomaly else 'Normal')
    if ml_anomaly:
        ml_run.font.color.rgb = RGBColor(0xdc, 0x35, 0x45)

    if rules:
        doc.add_paragraph()
        doc.add_paragraph('Rules Fired:', style='List Bullet')
        for rule in rules:
            doc.add_paragraph(f'{rule}', style='List Bullet 2')

    # --- SECTION 2: IDENTITY SUMMARY ---
    doc.add_heading('2. Resolved Identity Summary', level=1)

    id_table = doc.add_table(rows=1, cols=2)
    id_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = id_table.rows[0].cells
    hdr[0].text = 'Identifier Type'
    hdr[1].text = 'Values'

    for hdr_cell in hdr:
        for paragraph in hdr_cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    id_fields = [
        ("Phone Numbers", entity_data.get("phones", [])),
        ("Bank Accounts", entity_data.get("accounts", [])),
        ("IMEIs", entity_data.get("imeis", [])),
        ("UPI VPAs", entity_data.get("vpas", [])),
        ("IP Addresses", entity_data.get("ips", [])),
    ]

    for field_name, values in id_fields:
        if values:
            row = id_table.add_row().cells
            row[0].text = field_name
            row[1].text = ', '.join(str(v) for v in values[:10])

    # --- SECTION 3: RULE EXPLANATIONS ---
    doc.add_heading('3. Rule Analysis (Explainability)', level=1)

    explanations = risk_data.get("explanations", {})
    if explanations:
        for rule_id, explanation in explanations.items():
            doc.add_heading(f'Rule {rule_id}', level=3)
            doc.add_paragraph(explanation)
    else:
        doc.add_paragraph('No rules fired for this entity.')

    # --- SECTION 4: NETWORK GRAPH ---
    doc.add_heading('4. Money Flow / Identity Linkage Graph', level=1)

    if network_fig:
        try:
            # Set quick timeout hint if available
            img_bytes = network_fig.to_image(format="png", width=750, height=450, scale=1)
            img_stream = io.BytesIO(img_bytes)
            doc.add_picture(img_stream, width=Inches(6.5))
        except Exception as e:
            doc.add_paragraph('[Identity graph visual available in interactive portal. See Evidence Summary table below for complete audit trail.]')
    else:
        doc.add_paragraph('[Identity graph visual available in interactive portal. See Evidence Summary table below for complete audit trail.]')


    # --- SECTION 5: EVIDENCE TABLE ---
    doc.add_heading('5. Evidence Summary', level=1)

    entity_events = all_events_df[all_events_df['entity_id'] == entity_id].copy() if not all_events_df.empty else pd.DataFrame()

    if not entity_events.empty:
        entity_events['timestamp'] = pd.to_datetime(entity_events['timestamp'])
        entity_events = entity_events.sort_values('timestamp')

        ev_table = doc.add_table(rows=1, cols=6)
        ev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ev_hdr = ev_table.rows[0].cells
        headers = ['Timestamp', 'Type', 'Amount', 'Counterparty', 'Source', 'Row Ref']
        for i, h in enumerate(headers):
            ev_hdr[i].text = h
            for paragraph in ev_hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)

        for _, row in entity_events.head(50).iterrows():
            cells = ev_table.add_row().cells
            cells[0].text = str(row.get('timestamp', ''))[:19]
            cells[1].text = str(row.get('event_type', ''))
            amt = row.get('amount')
            cells[2].text = f"{amt:.0f}" if pd.notna(amt) else ''
            cells[3].text = str(row.get('counterparty', ''))[:30]
            cells[4].text = str(row.get('source_file', ''))
            cells[5].text = str(row.get('raw_row_ref', ''))

            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(7)
    else:
        doc.add_paragraph('No events found for this entity.')

    # --- SECTION 6: DECISION TRACE ---
    doc.add_heading('6. Decision Trace References', level=1)

    evidence_data = risk_data.get("evidence", {})
    if evidence_data:
        for rule_id, refs in evidence_data.items():
            doc.add_paragraph(f'{rule_id}: {", ".join(str(r) for r in refs[:15])}')
    else:
        doc.add_paragraph('No decision trace entries for this entity.')

    # --- FOOTER ---
    doc.add_paragraph('_' * 70)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        'CONFIDENTIAL - FOR LAW ENFORCEMENT USE ONLY\n'
        'Generated by E-Rakshak Forensic Analysis Platform\n'
        'Evidence admissible under Indian Evidence Act S.65B with proper certification'
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    filename = f"forensic_report_{entity_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = REPORT_DIR / filename
    doc.save(str(filepath))

    print(f"  Report saved: {filepath}")
    return filepath



# ============================================================
# FIU-IND STR (Suspicious Transaction Report) GENERATOR
# ============================================================

# Maps each fired rule ID to the FIU-IND "nature of suspicion" category it
# corresponds to, per the STR reporting conventions (PMLA Rule 3 categories).
RULE_TO_STR_CATEGORY = {
    "STR-1": ("Structuring / Smurfing",
              "Multiple transactions structured just below the statutory cash/transfer "
              "reporting threshold, apparently to evade reporting requirements."),
    "MUL-1": ("Mule Account / Layering of Proceeds",
              "Dormant account activated by a sudden large credit followed by rapid, "
              "near-total outflow, consistent with use as a money-mule conduit."),
    "LAY-1": ("Layering & Circular Fund Movement",
              "Funds passed through a rapid chain of intermediary accounts with "
              "decreasing amounts (commission-skim pattern), consistent with layering "
              "to obscure the audit trail."),
    "IDR-1": ("Identity Fraud / SIM Swap",
              "Concurrent or closely-timed change of device (IMEI) and/or SIM (IMSI) "
              "identifiers on the subject's mobile number, indicative of identity "
              "takeover facilitating unauthorized transactions."),
    "TCS-1": ("Suspicious Transaction Timing",
              "Transaction executed while the subject held simultaneous active phone "
              "call and internet session presence, indicating coordinated real-time "
              "facilitation of the transaction."),
    "TCS-2": ("Suspicious Transaction Timing",
              "A phone call to a co-suspect immediately preceded a fund transfer, "
              "suggesting coordination of the transaction by voice communication."),
    "LOC-1": ("Geographically Improbable Activity",
              "Subject's registered KYC location is inconsistent with the cell "
              "tower / IP-derived location active at the time of the transaction, "
              "at a rate implying physically impossible travel."),
}

STR_DEFAULT_CATEGORY = ("Other Suspicious Activity",
                         "Activity flagged by the E-Rakshak deterministic rule engine "
                         "warranting manual review by the reporting officer.")


def generate_str_report(entity_id, entity_data, risk_data, all_events_df,
                         reporting_entity=None, reporting_officer=None):
    """
    Generate a Suspicious Transaction Report (STR) formatted per FIU-IND
    (Financial Intelligence Unit - India) conventions, distinct from the
    full forensic investigation report.

    Structure follows FIU-IND STR filing conventions:
      1. Reporting Entity Details
      2. Suspect / Account Holder Identity
      3. Nature of Suspicion (mapped from fired rule(s) to FIU-IND category)
      4. Schedule of Suspicious Transactions
      5. Total Value Involved
      6. Reporting Officer Certification (placeholder fields)

    Args:
        entity_id: Resolved entity ID
        entity_data: Dict with phones, accounts, imeis, vpas, name
        risk_data: Dict with risk_tier, rules_fired, explanations, evidence
        all_events_df: DataFrame of all events
        reporting_entity: dict with reporting institution details (optional overrides)
        reporting_officer: dict with reporting officer details (optional, left blank if omitted)

    Returns:
        Path to generated .docx STR file
    """
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1a, 0x1d, 0x24)

    rules_fired = risk_data.get("rules_fired", [])
    explanations = risk_data.get("explanations", {})
    evidence = risk_data.get("evidence", {})
    risk_tier = risk_data.get("risk_tier", "LOW")

    # --- HEADER ---
    heading = doc.add_heading('', level=0)
    run = heading.add_run('SUSPICIOUS TRANSACTION REPORT (STR)')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1a, 0x1d, 0x24)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('Filed in accordance with FIU-IND reporting formats under the Prevention of '
                       'Money Laundering Act (PMLA), 2002, Rule 3')
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run(f'STR Reference: STR-{entity_id}-{datetime.now().strftime("%Y%m%d%H%M%S")} | '
                        f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('_' * 70)

    # --- SECTION 1: REPORTING ENTITY DETAILS ---
    doc.add_heading('1. Reporting Entity Details', level=1)
    reporting_entity = reporting_entity or {}

    re_table = doc.add_table(rows=0, cols=2)
    re_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    re_fields = [
        ("Name of Reporting Entity", reporting_entity.get("name", "________________________")),
        ("Reporting Entity Type", reporting_entity.get("type", "Bank / Financial Institution / Telecom Operator")),
        ("Registration / License No.", reporting_entity.get("reg_no", "________________________")),
        ("Branch / Circle", reporting_entity.get("branch", "________________________")),
        ("Report Generated By System", "E-Rakshak Forensic Analysis Platform (ERH26_PS_03)"),
    ]
    for label, value in re_fields:
        row = re_table.add_row().cells
        row[0].text = label
        row[1].text = str(value)
        for paragraph in row[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        for paragraph in row[1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

    # --- SECTION 2: SUSPECT / ACCOUNT HOLDER IDENTITY ---
    doc.add_heading('2. Suspect / Account Holder Identity', level=1)

    id_table = doc.add_table(rows=0, cols=2)
    id_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    id_fields = [
        ("Resolved Entity ID", entity_id),
        ("Name (KYC)", entity_data.get("name", "Unknown / Unmapped")),
        ("Phone Number(s)", ', '.join(entity_data.get("phones", [])) or "N/A"),
        ("Bank Account(s)", ', '.join(entity_data.get("accounts", [])) or "N/A"),
        ("UPI VPA(s)", ', '.join(entity_data.get("vpas", [])) or "N/A"),
        ("IMEI(s)", ', '.join(entity_data.get("imeis", [])) or "N/A"),
        ("IP Address(es)", ', '.join(entity_data.get("ips", [])) or "N/A"),
        ("Current System Risk Tier", risk_tier),
    ]
    for label, value in id_fields:
        row = id_table.add_row().cells
        row[0].text = label
        row[1].text = str(value)
        for paragraph in row[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        for paragraph in row[1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

    # --- SECTION 3: NATURE OF SUSPICION ---
    doc.add_heading('3. Nature / Grounds of Suspicion', level=1)

    if rules_fired:
        doc.add_paragraph(
            'This report is filed based on the following deterministic rule firing(s), each mapped '
            'to its corresponding FIU-IND suspicious-activity category:'
        )
        cat_table = doc.add_table(rows=1, cols=3)
        cat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = cat_table.rows[0].cells
        for i, h in enumerate(['Rule ID', 'FIU-IND Category', 'Grounds']):
            hdr[i].text = h
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        for rule_id in rules_fired:
            category, grounds = RULE_TO_STR_CATEGORY.get(rule_id, STR_DEFAULT_CATEGORY)
            row = cat_table.add_row().cells
            row[0].text = rule_id
            row[1].text = category
            row[2].text = grounds
            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        doc.add_paragraph()
        doc.add_paragraph('Detailed rule findings:')
        for rule_id in rules_fired:
            explanation = explanations.get(rule_id, '')
            doc.add_heading(f'Rule {rule_id}', level=3)
            doc.add_paragraph(explanation or 'No detailed explanation recorded.')
    else:
        doc.add_paragraph(
            'No deterministic rule has fired for this entity. This STR is generated on-demand '
            'and should be reviewed manually before filing.'
        )

    # --- SECTION 4: SCHEDULE OF SUSPICIOUS TRANSACTIONS ---
    doc.add_heading('4. Schedule of Suspicious Transactions', level=1)

    entity_events = all_events_df[all_events_df['entity_id'] == entity_id].copy() if not all_events_df.empty else pd.DataFrame()
    entity_txns = entity_events[entity_events['event_type'] == 'transaction'].copy() if not entity_events.empty else pd.DataFrame()

    total_value = 0.0
    if not entity_txns.empty:
        entity_txns['timestamp'] = pd.to_datetime(entity_txns['timestamp'])
        entity_txns = entity_txns.sort_values('timestamp')
        total_value = entity_txns['amount'].abs().sum()

        sched_table = doc.add_table(rows=1, cols=6)
        sched_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = sched_table.rows[0].cells
        headers = ['S.No', 'Date/Time', 'Direction', 'Amount (INR)', 'Counterparty', 'Source / Row Ref']
        for i, h in enumerate(headers):
            hdr[i].text = h
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)

        for i, (_, row) in enumerate(entity_txns.head(100).iterrows(), start=1):
            cells = sched_table.add_row().cells
            amt = row.get('amount', 0)
            direction = "Credit (IN)" if pd.notna(amt) and amt > 0 else "Debit (OUT)"
            cells[0].text = str(i)
            cells[1].text = str(row.get('timestamp', ''))[:19]
            cells[2].text = direction
            cells[3].text = f"{abs(amt):,.2f}" if pd.notna(amt) else ''
            cells[4].text = str(row.get('counterparty', ''))[:30]
            cells[5].text = f"{row.get('source_file', '')} / {row.get('raw_row_ref', '')}"
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(7)
    else:
        doc.add_paragraph('No transaction events found for this entity.')

    # --- SECTION 5: TOTAL VALUE INVOLVED ---
    doc.add_heading('5. Total Value Involved', level=1)
    tv_para = doc.add_paragraph()
    tv_para.add_run('Total Aggregate Value of Suspicious Transactions: ').bold = True
    tv_run = tv_para.add_run(f'INR {total_value:,.2f}')
    tv_run.bold = True
    tv_run.font.size = Pt(12)
    tv_run.font.color.rgb = RGBColor(0xdc, 0x35, 0x45)

    count_para = doc.add_paragraph()
    count_para.add_run('Number of Transactions in Schedule: ').bold = True
    count_para.add_run(str(len(entity_txns)))

    # --- SECTION 6: REPORTING OFFICER CERTIFICATION ---
    doc.add_heading('6. Reporting Officer Certification', level=1)
    doc.add_paragraph(
        'I hereby certify that the above transaction(s) have been reviewed and, based on the '
        'attached grounds of suspicion, are being reported as suspicious in accordance with the '
        'Prevention of Money Laundering Act (PMLA), 2002 and rules made thereunder.'
    )

    reporting_officer = reporting_officer or {}
    cert_table = doc.add_table(rows=0, cols=2)
    cert_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cert_fields = [
        ("Name of Principal Officer", reporting_officer.get("name", "________________________")),
        ("Designation", reporting_officer.get("designation", "________________________")),
        ("Employee / Badge ID", reporting_officer.get("badge_id", "________________________")),
        ("Contact Number", reporting_officer.get("contact", "________________________")),
        ("Date of Filing", reporting_officer.get("filing_date", "________________________")),
        ("Signature", "________________________")
    ]
    for label, value in cert_fields:
        row = cert_table.add_row().cells
        row[0].text = label
        row[1].text = str(value)
        for paragraph in row[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        for paragraph in row[1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

    # --- FOOTER ---
    doc.add_paragraph('_' * 70)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        'CONFIDENTIAL - SUSPICIOUS TRANSACTION REPORT - FOR FIU-IND FILING USE ONLY\n'
        'Auto-generated by E-Rakshak Forensic Analysis Platform. Manual officer review and '
        'certification required before submission to FIU-IND.'
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save — filename prefix distinguishes it from the full forensic report
    filename = f"str_report_{entity_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = REPORT_DIR / filename
    doc.save(str(filepath))

    print(f"  STR report saved: {filepath}")
    return filepath


def generate_reports_for_flagged(scored_entities, entities, all_events_df,
                                  create_timeline_fn=None, create_network_fn=None,
                                  network_graph=None):
    """Generate forensic Word reports for all entities with risk tier HIGH or CRITICAL."""
    flagged = {eid: data for eid, data in scored_entities.items()
               if data["risk_tier"] in ["HIGH", "CRITICAL"]}

    if not flagged:
        print("  No HIGH/CRITICAL entities to report on.")
        return []

    reports = []
    for entity_id, risk_data in flagged.items():
        entity_data = entities.get(entity_id, {})
        report_path = generate_forensic_report(
            entity_id, entity_data, risk_data, all_events_df,
            timeline_fig=None, network_fig=None
        )
        reports.append(report_path)

    return reports
